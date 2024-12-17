import requests
import os
import re
import pickle
import nltk
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document
nltk.download('punkt_tab')



#URL of the book from project gutenberg
url = "https://www.gutenberg.org/cache/epub/100/pg100.txt"
output_dir = "./ai_search_phase_IV"
f_type = None

def read_document():
    url_path = None
    user_input = input("Please enter a URL or file path: ")

    # Check if the input is a URL
    if user_input.startswith(('http://', 'https://')):
        # Determine file type based on URL extension
        file_type = user_input[-4:].lower()
        response = requests.get(user_input)
        response.raise_for_status()

        if file_type == '.txt':
            url_path = response.text
        elif file_type == '.pdf':
            pdf_reader = PdfReader(BytesIO(response.content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            url_path = text
        elif file_type == 'docx':
            doc = Document(BytesIO(response.content))
            text = " ".join([paragraph.text for paragraph in doc.paragraphs])
            url_path = text
        else:
            print("Unsupported file type via URL.")
    
    # If input is not a URL, assume it's a file path
    else:
        if os.path.exists(user_input):
            file_type = user_input[-4:].lower()

            if file_type == '.txt':
                with open(user_input, 'r') as file:
                    url_path = file.read()
            elif file_type == '.pdf':
                with open(user_input, 'rb') as file:
                    pdf_reader = PdfReader(BytesIO(file.read()))
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()
                    url_path = text
            elif file_type == 'docx':
                with open(user_input, 'rb') as file:
                    doc = Document(BytesIO(file.read()))
                    text = " ".join([paragraph.text for paragraph in doc.paragraphs])
                    url_path = text
            else:
                print("Unsupported file type.")
        else:
            print("The provided file path does not exist.")
    
    return url_path, file_type, user_input

def download_file(url : str, file_name : str) -> str:
    #send a get request
    response = requests.get(url)

    # Check if the request was successful (status code 200)
    if response.status_code == 200:
        # Save the content to a .txt file
        
        # Create data directory if it doesn't exist
        os.makedirs(os.path.join(output_dir, "data"), exist_ok=True)
        data_file = os.path.join(output_dir, "data", "text")
        
        with open(data_file, "w", encoding='utf-8') as file:
            file.write(response.text)
            print("Book downloaded successfully!")
    else:
        print(f"Failed to download the book. Status code: {response.status_code}")

    #call the utility methods to cleanup the file 
    #read between markers:
    content = read_between_markers(data_file)

    #write between markers
    file_obj = write_between_markers(data_file, content)

    return file_obj

'''
* Function to read the file from project gutenburg within the markers.
* This way we get just the actual content and not licensing stuff.
# Gutenburg doc starts with ***start and ends with ***end
'''
def read_between_markers(file_path : str) -> list:
    # List containing the content between the markers.
    content = []

    #open the file
    with open (file_path, 'r') as file:
        start_reading = False
        #iterate through the lines.
        for line in file:
            if "*** START" in line:
                start_reading = True
                continue
            if "*** END" in line:
                start_reading = False
                break
            if start_reading:
                content.append(line.strip())
        return content
    
'''
Function that takes in the content and writes between the markers.
'''
def write_between_markers(file_path : str, content : list) -> str:
    with open(file_path, 'w', encoding='utf-8') as file:
        for line in content:
            file.write(line + "\n")
    
    return file_path

'''
* File that converts the file into chunks of paragraph.
'''
def convert_to_chunks(text : str) -> list:    
    # Tokenize paragraphs using regular expressions
    if f_type == '.txt':
        paragraphs = re.split(r'\n\s*\n', text)
    elif f_type == 'docx':
        paragraphs = text.split('\n')
    elif f_type == '.pdf':
        new_text = text.replace('\n', ' ')
        paragraphs = nltk.sent_tokenize(new_text)

    return paragraphs

'''
*   This method is used to write vector embeddings to a pickle file.
*   This will save as a cache and avoid computing multiple times.
'''
def save_to_pickle(vector_embeddings):
    # Save embeddings
    # Create data directory if it doesn't exist
    os.makedirs(os.path.join(output_dir, "data"), exist_ok=True)
    data_file = os.path.join(output_dir, "data", "embeddings.pkl")
    
    #save the pickle file in data    
    with open(data_file, 'wb') as f:
        pickle.dump(vector_embeddings, f)

'''
* This function will serve as a cache
'''
def read_from_pickle(file_obj : str) -> list:

    with open(file_obj, 'rb') as f:
        vector_embeddings = pickle.load(f)

    return  vector_embeddings



if __name__ == "__main__":
    file_obj = download_file(url, "complete_works_of_william_shakespeare.txt")
    res_list = convert_to_chunks(file_obj)

    print(f"{len(res_list)}")